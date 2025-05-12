from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_modern_clean_template(output_path, user_data):
    doc = Document()

    # Full Name (Centered, Big)
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(user_data.get("Full Name", ""))
    run.font.size = Pt(18)
    run.bold = True

    # Contact Info (Links in Blue and Comma-separated, Centered)
    contact_parts = []
    for field in ["Email", "LinkedIn", "GitHub"]:
        value = user_data.get(field, "")
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, value in enumerate(contact_parts):
            # Add a link-style for email, LinkedIn, and GitHub
            run = contact.add_run(value)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
            run.underline = True
            # Add a comma separator except for the last link
            if i < len(contact_parts) - 1:
                contact.add_run(", ")

        doc.add_paragraph()  # spacing

    # Resume Sections
    add_section(doc, "Professional Summary", user_data.get("Professional Summary", ""))

    add_multientry_section(doc, "Projects", user_data.get("Projects", []), [
        ("Project Title", "title"),
        ("Project Description", "description"),
        ("Technologies Used", "techStack"),
        ("Project URL", "link")
    ])

    add_multientry_section(doc, "Work Experience", user_data.get("Experiences", []), [
        ("Company Name", "companyName"),
        ("Job Title", "jobTitle"),
        ("Duration", "duration"),
        ("Job Responsibilities", "description")
    ])

    add_multientry_section(doc, "Certifications", user_data.get("Certifications", []), [
        ("Certification Title", "title"),
        ("Issuer", "issuer"),
        ("Date", "date")
    ])

    add_multientry_section(doc, "Education", user_data.get("Education", []), [
        ("Degree", "degree"),
        ("Institution Name", "institution"),
        ("Duration", "duration")
    ])

    # Skills (comma-separated)
    add_skill_section(doc, "Technical Skills", user_data.get("TechnicalSkills", []))
    add_skill_section(doc, "Soft Skills", user_data.get("SoftSkills", []))

    doc.save(output_path)


def add_section(doc, title, content):
    if not content.strip():
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    para = doc.add_paragraph()
    para.add_run(content).font.size = Pt(12)


def add_multientry_section(doc, title, entries, fields):
    if not entries:
        return

    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    for entry in entries:
        for label, key in fields:
            value = entry.get(key, "")
            if value:
                bullet_para = doc.add_paragraph(style="List Bullet")
                label_run = bullet_para.add_run(f"{label}: ")
                label_run.font.size = Pt(12)
                label_run.bold = True

                value_run = bullet_para.add_run(str(value))
                value_run.font.size = Pt(12)

        doc.add_paragraph()  # Space between entries


def add_skill_section(doc, title, skills):
    if not skills:
        return

    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    skill_text = ", ".join(skills)
    para = doc.add_paragraph()
    para.add_run(skill_text).font.size = Pt(12)


def generate_resume_from_form_data(user_data, output_path):
    """
    Transforms user data and calls create_modern_clean_template to generate the resume document.
    """
    transformed_data = {
        "Full Name": user_data.get("Full Name", ""),
        "Email": user_data.get("Email", ""),
        "LinkedIn": user_data.get("LinkedIn", ""),
        "GitHub": user_data.get("GitHub", ""),
        "Professional Summary": user_data.get("Professional Summary", ""),
        "Projects": user_data.get("Projects", []),
        "Experiences": user_data.get("Experiences", []) or user_data.get("Experience", []),
        "Certifications": user_data.get("Certifications", []) or user_data.get("Certification", []),
        "Education": user_data.get("Education", []),
        "TechnicalSkills": user_data.get("TechnicalSkills", []) or user_data.get("technicalSkills", []),
        "SoftSkills": user_data.get("SoftSkills", []) or user_data.get("softSkills", []),
    }

    create_modern_clean_template(output_path, transformed_data)

