from flask import Flask, request, jsonify, send_file
from flask_cors import CORS, cross_origin
import psycopg2
import bcrypt
import os
import traceback
import json

# Import the resume generation code
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_modern_clean_template(output_path, user_data):
    doc = Document()

    # Full Name (Centered, Big)
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(user_data.get("Full Name", ""))
    run.font.size = Pt(18)
    run.bold = True

    # Contact Info (Links in Blue and Comma-separated, Centered)
    contact_parts = []
    for field in ["Email", "LinkedIn", "GitHub"]:
        value = user_data.get(field, "")
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, value in enumerate(contact_parts):
            run = contact.add_run(value)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
            if i < len(contact_parts) - 1:
                contact.add_run(", ")
        doc.add_paragraph()

    # Resume Sections
    add_section(doc, "Professional Summary", user_data.get("Professional Summary", ""))

    add_multientry_section(doc, "Projects", user_data.get("Projects", []), [
        ("Project Title", "title"),
        ("Project Description", "description"),
        ("Technologies Used", "techStack"),
        ("Project URL", "link")
    ])

    add_multientry_section(doc, "Work Experience", user_data.get("Experiences", []), [
        ("Company Name", "companyName"),
        ("Job Title", "jobTitle"),
        ("Duration", "duration"),
        ("Job Responsibilities", "description")
    ])

    add_multientry_section(doc, "Certifications", user_data.get("Certifications", []), [
        ("Certification Title", "title"),
        ("Issuer", "issuer"),
        ("Date", "date")
    ])

    add_multientry_section(doc, "Education", user_data.get("Education", []), [
        ("Degree", "degree"),
        ("Institution Name", "institution"),
        ("Duration", "duration")
    ])

    # Skills (comma-separated)
    add_skill_section(doc, "Technical Skills", user_data.get("TechnicalSkills", []))
    add_skill_section(doc, "Soft Skills", user_data.get("SoftSkills", []))

    doc.save(output_path)


def add_section(doc, title, content):
    if not content or not content.strip():
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    para = doc.add_paragraph()
    para.add_run(content).font.size = Pt(12)


def add_multientry_section(doc, title, entries, fields):
    if not entries:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    for entry in entries:
        for label, key in fields:
            value = entry.get(key, "")
            if value:
                bullet_para = doc.add_paragraph(style="List Bullet")
                label_run = bullet_para.add_run(f"{label}: ")
                label_run.font.size = Pt(12)
                label_run.bold = True

                value_run = bullet_para.add_run(str(value))
                value_run.font.size = Pt(12)
        doc.add_paragraph()


def add_skill_section(doc, title, skills):
    if not skills:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.underline = True

    skill_text = ", ".join(skills)
    para = doc.add_paragraph()
    para.add_run(skill_text).font.size = Pt(12)


# Initialize Flask app
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Database connection
try:
    conn = psycopg2.connect(
        dbname="resumeDB",
        user="postgres",
        password="112255",
        host="localhost",
        port="5432"
    )
    conn.autocommit = True
    cur = conn.cursor()
    print("✅ Database connected successfully!")
except Exception as e:
    print("❌ Failed to connect to database:", e)
    conn = None
    cur = None
    
    
# ---------------- login ROUTE ----------------
@app.route('/login', methods=['POST', 'OPTIONS'])
@cross_origin()
def login():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')

        if not username or not password:
            return jsonify({'message': 'Username and password are required'}), 400

        cur.execute("SELECT password FROM users WHERE username = %s", (username,))
        result = cur.fetchone()

        if not result:
            return jsonify({'message': 'User not found'}), 404

        hashed_pw = result[0]

        if bcrypt.checkpw(password.encode('utf-8'), hashed_pw.encode('utf-8')):
            return jsonify({'message': 'Login successful'}), 200
        else:
            return jsonify({'message': 'Incorrect password'}), 401

    except Exception as e:
        print(f"❌ Login error: {e}")
        traceback.print_exc()
        return jsonify({'message': f'Login failed: {str(e)}'}), 500


# ---------------- SIGNUP ROUTE ----------------
@app.route('/signup', methods=['POST', 'OPTIONS'])
@cross_origin()
def signup():
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        data = request.get_json()
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')

        if not username or not email or not password:
            return jsonify({'message': 'All fields are required'}), 400

        hashed_pw = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

        cur.execute(
            "INSERT INTO users (username, email, password) VALUES (%s, %s, %s)",
            (username, email, hashed_pw)
        )

        return jsonify({'message': 'User registered successfully'}), 201

    except Exception as e:
        print(f"❌ Signup error: {e}")
        traceback.print_exc()
        return jsonify({'message': f'Signup failed: {str(e)}'}), 500

# ---------------- RESUME GENERATION ----------------
@app.route('/generate_resume', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_resume():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        user_data = request.get_json()
        if not user_data:
            return jsonify({'message': 'No user data provided'}), 400

        print("Received data for resume generation:")
        print(json.dumps(user_data, indent=2))

        output_path = os.path.join(os.getcwd(), 'generated_resume.docx')
        create_modern_clean_template(output_path, user_data)

        print(f"✅ Resume generated at {output_path}")
        return send_file(
            output_path,
            as_attachment=True,
            download_name='resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"❌ Error generating resume: {e}")
        traceback.print_exc()
        return jsonify({'message': f'Error generating resume: {e}'}), 500

@app.route("/", methods=["GET", "HEAD"])
def home():
    return "Resume Builder API is running!", 200
# ---------------- SERVER START ----------------
if __name__ == '__main__':
    print("Starting Resume Builder API on http://localhost:3001")
    app.run(host='0.0.0.0', port=3001, debug=True)